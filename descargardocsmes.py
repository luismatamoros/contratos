import streamlit as st
import os
import tempfile
import zipfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageChops
import datetime

# --- CONFIGURACIÓN GENERAL ---
CARPETA_DOCUMENTOS = "./documentos"
CARPETA_ADJUNTOS = "./adjuntos"
ANCHO_MAXIMO = Inches(6.0)

def textos_por_mes(nombre_mes):
    año = datetime.datetime.now().year
    textoPeriodo = f"{nombre_mes}, {año}"
    textoRige = f"1ero de {nombre_mes} a 30 de {nombre_mes}, {año}"
    textoDurante = f"durante el mes de {nombre_mes.lower()}"
    textoEliminar = "No se presenta ninguna incapacidad..."
    return {
        "textoPeriodo": textoPeriodo,
        "textoRigeAPartirDe": textoRige,
        "textoDuranteElMes": textoDurante
    }, textoEliminar

def recortar_bordes_blancos(imagen):
    imagen = imagen.convert("RGB")
    fondo = Image.new("RGB", imagen.size, (255, 255, 255))
    diferencia = ImageChops.difference(imagen, fondo)
    bbox = diferencia.getbbox()
    return imagen.crop(bbox) if bbox else imagen

def contiene_anexos(parrafos):
    for par in parrafos[-5:]:
        if "anexos" in par.text.strip().lower():
            return True
    return False

def agregar_anexos(doc):
    imagenes = [f for f in os.listdir(CARPETA_ADJUNTOS) if f.lower().endswith((".png", ".jpg", ".jpeg"))]
    if not imagenes:
        return  # No hay imágenes para anexar

    # Añadir página nueva si no hay sección de anexos
    if not contiene_anexos(doc.paragraphs):
        doc.add_page_break()
        par_anexos = doc.add_paragraph("Anexos")
        par_anexos.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par_anexos.runs[0]
        run.font.size = Pt(14)
        run.bold = True

    for img_name in imagenes:
        img_path = os.path.join(CARPETA_ADJUNTOS, img_name)
        img = recortar_bordes_blancos(Image.open(img_path))
        temp_img_path = os.path.join(tempfile.gettempdir(), img_name)
        img.save(temp_img_path)

        par_img = doc.add_paragraph(f"Anexo: {img_name}")
        par_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = par_img.runs[0]
        run_img.font.size = Pt(12)

        doc.add_picture(temp_img_path, width=ANCHO_MAXIMO)
        doc.add_paragraph()  # Espacio extra

def procesar_documento(doc_path, reemplazos, texto_a_eliminar):
    doc = Document(doc_path)

    # Reemplazos y eliminaciones en párrafos
    for par in doc.paragraphs:
        for buscar, reemplazo in reemplazos.items():
            if buscar in par.text:
                for run in par.runs:
                    run.text = run.text.replace(buscar, reemplazo)
        if texto_a_eliminar in par.text:
            for run in par.runs:
                run.text = run.text.replace(texto_a_eliminar, "")

    # Reemplazos y eliminaciones en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for par in celda.paragraphs:
                    for buscar, reemplazo in reemplazos.items():
                        if buscar in par.text:
                            for run in par.runs:
                                run.text = run.text.replace(buscar, reemplazo)
                    if texto_a_eliminar in par.text:
                        for run in par.runs:
                            run.text = run.text.replace(texto_a_eliminar, "")

    # Reemplazos y eliminaciones en header y footer
    for section in doc.sections:
        for parte in [section.header, section.footer]:
            for par in parte.paragraphs:
                for buscar, reemplazo in reemplazos.items():
                    if buscar in par.text:
                        for run in par.runs:
                            run.text = run.text.replace(buscar, reemplazo)
                if texto_a_eliminar in par.text:
                    for run in par.runs:
                        run.text = run.text.replace(texto_a_eliminar, "")

    # Agregar anexos de imágenes de la carpeta adjuntos
    agregar_anexos(doc)

    output_path = os.path.join(tempfile.gettempdir(), f"{os.path.basename(doc_path)}")
    doc.save(output_path)
    return output_path

# --- INTERFAZ STREAMLIT ---
st.title("📄 Generar y descargar TODOS los documentos modificados con anexos de un mes")

mes = st.selectbox("Selecciona el mes", [
    "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"
])

if st.button("🚀 Procesar y descargar todos los documentos"):
    reemplazos, texto_eliminar = textos_por_mes(mes)
    archivos_generados = []

    for archivo in os.listdir(CARPETA_DOCUMENTOS):
        if archivo.endswith(".docx"):
            ruta_doc = os.path.join(CARPETA_DOCUMENTOS, archivo)
            salida = procesar_documento(ruta_doc, reemplazos, texto_eliminar)
            archivos_generados.append(salida)

    if archivos_generados:
        zip_path = os.path.join(tempfile.gettempdir(), f"documentos_{mes}.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for fpath in archivos_generados:
                zipf.write(fpath, arcname=os.path.basename(fpath))

        with open(zip_path, "rb") as fzip:
            st.download_button(f"📥 Descargar ZIP con documentos modificados y anexos de {mes}", fzip, file_name=f"documentos_{mes}.zip")
    else:
        st.warning("No se encontraron documentos para procesar.")