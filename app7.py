import streamlit as st
import os
import tempfile
from PIL import Image, ImageChops
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import json

# --- CONFIGURACIÓN GENERAL ---
CARPETA_DOCUMENTOS = "./documentos"
ARCHIVO_ESTADO = "estado_imagenes.json"
extensiones_imagenes = [".png", ".jpg", ".jpeg"]
ancho_maximo = Inches(6.0)

# --- FUNCIONES ---
def cargar_estado():
    if os.path.exists(ARCHIVO_ESTADO):
        with open(ARCHIVO_ESTADO, "r") as f:
            return json.load(f)
    return {}

def guardar_estado(estado):
    with open(ARCHIVO_ESTADO, "w") as f:
        json.dump(estado, f, indent=4)

def textos_por_mes(nombre_mes):
    mes_num = {
        "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
        "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
    }
    numero = mes_num[nombre_mes]
    año = datetime.datetime.now().year
    textoPeriodo = f"{nombre_mes}, {año}"
    textoRige = f"1ero de {nombre_mes} a 30 de {nombre_mes}, {año}"
    textoDurante = f"durante el mes de {nombre_mes.lower()}"
    textoEliminar = "No se presenta ninguna incapacidad..."
    return {
        "textoPeriodo": textoPeriodo,
        "textoRigeAPartirDe": textoRige,
        "textoDuranteElMes": textoDurante
    }, textoEliminar

def recortar_bordes_blancos(imagen):
    imagen = imagen.convert("RGB")
    fondo = Image.new("RGB", imagen.size, (255, 255, 255))
    diferencia = ImageChops.difference(imagen, fondo)
    bbox = diferencia.getbbox()
    return imagen.crop(bbox) if bbox else imagen

def contiene_anexos(parrafos):
    for par in parrafos[-5:]:
        if "anexos" in par.text.strip().lower():
            return True
    return False

def procesar_documento(doc_path, imagen, reemplazos, texto_a_eliminar):
    doc = Document(doc_path)

    for par in doc.paragraphs:
        for buscar, reemplazo in reemplazos.items():
            if buscar in par.text:
                for run in par.runs:
                    run.text = run.text.replace(buscar, reemplazo)
        if texto_a_eliminar in par.text:
            for run in par.runs:
                run.text = run.text.replace(texto_a_eliminar, "")

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for par in celda.paragraphs:
                    for buscar, reemplazo in reemplazos.items():
                        if buscar in par.text:
                            for run in par.runs:
                                run.text = run.text.replace(buscar, reemplazo)
                    if texto_a_eliminar in par.text:
                        for run in par.runs:
                            run.text = run.text.replace(texto_a_eliminar, "")

    for section in doc.sections:
        for parte in [section.header, section.footer]:
            for par in parte.paragraphs:
                for buscar, reemplazo in reemplazos.items():
                    if buscar in par.text:
                        for run in par.runs:
                            run.text = run.text.replace(buscar, reemplazo)
                if texto_a_eliminar in par.text:
                    for run in par.runs:
                        run.text = run.text.replace(texto_a_eliminar, "")

    if imagen:
        if not contiene_anexos(doc.paragraphs):
            doc.add_page_break()
            par_anexos = doc.add_paragraph("Anexos")
            par_anexos.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par_anexos.runs[0]
            run.font.size = Pt(14)
            run.bold = True

        img = recortar_bordes_blancos(Image.open(imagen))
        w, h = img.size
        proporcion = ancho_maximo / w
        temp_img_path = os.path.join(tempfile.gettempdir(), "imagen_anexo.jpg")
        img.save(temp_img_path)
        par_img = doc.add_paragraph(f"Anexo: {os.path.basename(imagen.name)}")
        par_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = par_img.runs[0]
        run_img.font.size = Pt(12)
        doc.add_picture(temp_img_path, width=ancho_maximo)
        doc.add_paragraph()

    output_path = os.path.join(tempfile.gettempdir(), "documento_modificado.docx")
    doc.save(output_path)
    return output_path

# --- INTERFAZ STREAMLIT ---
st.title("📄 Generador de Documentos con Imagen Anexa")

# Cargar estado
estado = cargar_estado()

mes = st.selectbox("📆 Selecciona el mes", [
    "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"
])

# Documentos con indicador de estado
docs_disponibles = [
    f"{doc} ✅" if estado.get(doc) else doc
    for doc in os.listdir(CARPETA_DOCUMENTOS)
    if doc.endswith(".docx")
]
doc_opcion = st.selectbox("📂 Selecciona un documento", docs_disponibles)
doc_seleccionado = doc_opcion.replace(" ✅", "")

imagen_subida = st.file_uploader("🖼️ Sube una imagen (PNG, JPG)", type=["png", "jpg", "jpeg"])

if st.button("🚀 Generar documento"):
    if doc_seleccionado:
        doc_path = os.path.join(CARPETA_DOCUMENTOS, doc_seleccionado)
        reemplazos, texto_eliminar = textos_por_mes(mes)
        salida = procesar_documento(doc_path, imagen_subida, reemplazos, texto_eliminar)

        # Guardar estado de que se subió imagen
        if imagen_subida:
            estado[doc_seleccionado] = True
            guardar_estado(estado)

        with open(salida, "rb") as f:
            st.download_button("📥 Descargar documento modificado", f, file_name=f"{doc_seleccionado[:-5]}_modificado.docx")
    else:
        st.warning("Selecciona un documento.")